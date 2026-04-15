"""Quotation DOCX generator.

Builds a professional quotation document matching the company template.
All dimensions, fonts, and styles are derived from the reference template.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Emu, Pt, RGBColor
from docx.table import _Cell

_CONTENT_WIDTH = Cm(21.00) - Cm(1.91) - Cm(2.11)  # page minus margins

_TEMPLATES_DIR = Path(__file__).parent / "templates"
_IMAGES_DIR = _TEMPLATES_DIR / "images"

# Fonts from template analysis
_FONT_NAME = "Verdana"
_FONT_SIZE = Pt(10)
_FONT_COLOR = RGBColor(0, 0, 0)

# Page setup from template analysis
_PAGE_WIDTH = Cm(21.00)
_PAGE_HEIGHT = Cm(29.70)
_TOP_MARGIN = Cm(4.30)
_BOTTOM_MARGIN = Cm(2.38)
_LEFT_MARGIN = Cm(1.91)
_RIGHT_MARGIN = Cm(2.11)
_HEADER_DISTANCE = Cm(0.96)
_FOOTER_DISTANCE = Cm(2.03)

# Table column widths — description gets the bulk of space
_COL_MODEL = Cm(2.62)
_COL_DESC = Cm(9.82)
_COL_QTY = Cm(1.50)
_COL_UNIT = Cm(2.30)
_COL_TOTAL = Cm(2.50)


@dataclass
class QuotationProduct:
    code: str
    description: str
    quantity: int | float
    unit_price: float
    total_price: float


@dataclass
class QuotationData:
    client_name: str
    client_address: str
    reference_number: str
    generation_date: date
    project_name: str
    service_option: int  # 1, 2, or 3
    products: list[QuotationProduct]
    subtotal: float
    vat: float
    grand_total: float
    subject: str | None = None
    payment_terms_text: str | None = None
    device_count: int = 0
    installation_amount: float = 0.0
    inclusion_answers: dict[str, bool] | None = None
    letterhead_bytes: bytes | None = None
    signature_bytes: bytes | None = None
    signatory_name: str | None = None
    company_phone: str | None = None


def generate_quotation(data: QuotationData) -> bytes:
    """Generate a quotation DOCX and return it as bytes."""
    if data.letterhead_bytes:
        from app.shared.upload_security import (
            sanitize_docx_package,
            strip_docx_external_relationships,
        )

        # Strip macros/VBA/embedded objects from the ZIP, then open
        clean_bytes = sanitize_docx_package(data.letterhead_bytes)
        doc = Document(io.BytesIO(clean_bytes))
        # Remove any external URL references (tracking pixels etc.)
        strip_docx_external_relationships(doc)
        # Clear any body content from the letterhead template
        for p in doc.paragraphs:
            p.clear()
    else:
        # No letterhead uploaded — build from code with page setup only
        doc = Document()
        _setup_page(doc)
        # Only add header/footer images if they exist on disk
        if (_IMAGES_DIR / "logo.png").exists():
            _setup_header(doc)
        if (_IMAGES_DIR / "footer.png").exists():
            _setup_footer(doc)

    # Build body content — common top section
    _add_client_and_date(doc, data)
    _add_subject(doc, data.subject or f"Fire Alarm System \u2013 Simplex- {data.project_name}")
    _add_greeting(doc, data.client_name)

    if data.service_option == 1:
        # Supply-only template
        _add_intro(doc)
        _add_scope(doc, 1)
        _add_warranty(doc)
    else:
        # Installation template (options 2 and 3)
        _add_intro_installation(doc)
        _add_scope(doc, data.service_option)
        _add_body_text(doc,
            "Any changes in architectural design or material list or "
            "system design will change price."
        )
        _add_warranty_short(doc)

    # Common sections
    _add_cancellation(doc)
    _add_limitation_of_liability(doc)
    _add_notes_exclusions(doc, data)
    _add_prices_and_terms(doc)
    _add_payment_terms(doc, data)
    _add_time_for_supplies(doc)
    _add_validity_and_signature(doc, data)
    _add_product_table(doc, data)

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = _PAGE_WIDTH
    section.page_height = _PAGE_HEIGHT
    section.top_margin = _TOP_MARGIN
    section.bottom_margin = _BOTTOM_MARGIN
    section.left_margin = _LEFT_MARGIN
    section.right_margin = _RIGHT_MARGIN
    section.header_distance = _HEADER_DISTANCE
    section.footer_distance = _FOOTER_DISTANCE


def _setup_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Use anchored (floating) images matching reference template exactly
    p = header.paragraphs[0]
    run = p.add_run()
    header_part = header.part

    # Add logo as anchored image
    logo_path = str(_IMAGES_DIR / "logo.png")
    _add_anchored_image(
        run, logo_path,
        cx=984250, cy=1031875,       # 2.73cm × 2.87cm
        pos_h=68580, pos_v=-108585,   # 0.19cm right, 0.30cm above para
        part=header_part,
    )

    # Add header text as anchored image
    header_text_path = str(_IMAGES_DIR / "header_text.png")
    _add_anchored_image(
        run, header_text_path,
        cx=4443730, cy=737235,        # 12.34cm × 2.05cm
        pos_h=1188000, pos_v=220345,  # 3.30cm right, 0.61cm below para
        part=header_part,
    )

    # Give paragraph enough line height so border sits below the floating images
    # Lowest image bottom = 2.66cm from para; add buffer → 2.90cm ≈ 1644 twips
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "1644")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)

    # Thin bottom border — professional separator line below header images
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")       # 0.5pt thin line
    bottom.set(qn("w:space"), "1")    # minimal gap between line and spacing
    bottom.set(qn("w:color"), "A0A0A0")  # subtle gray
    pBdr.append(bottom)
    pPr.append(pBdr)


def _setup_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Anchored footer image — spans nearly full page width (edge to edge)
    p = footer.paragraphs[0]
    run = p.add_run()
    footer_path = str(_IMAGES_DIR / "footer.png")
    _add_anchored_image(
        run, footer_path,
        cx=7496175, cy=856615,         # 20.82cm × 2.38cm
        pos_h=-657225, pos_v=48260,    # -1.83cm left (into margin), 0.13cm below para
        part=footer.part,
    )


# ---------------------------------------------------------------------------
# Body content builders
# ---------------------------------------------------------------------------

def _add_client_and_date(doc: Document, data: QuotationData) -> None:
    d = data.generation_date
    day_suffix = _ordinal_suffix(d.day)
    ref_str = f"Ref.: MI/203-C/{data.reference_number}"

    # Right-tab stop at content width for right-aligned text
    tab_pos = _CONTENT_WIDTH

    # Line 1: client name [TAB] date
    p1 = doc.add_paragraph()
    _add_right_tab_stop(p1, tab_pos)
    run = p1.add_run(f"Engr, {data.client_name}")
    _style_run(run, bold=True)
    p1.add_run("\t")  # tab to right side
    run_d1 = p1.add_run(f"Date: {d.day}")
    _style_run(run_d1, bold=True)
    run_sup = p1.add_run(day_suffix)
    _style_run(run_sup, bold=True, size=Pt(7))
    run_sup.font.superscript = True
    run_d2 = p1.add_run(f" {d.strftime('%b')} {d.year}")
    _style_run(run_d2, bold=True)

    # Line 2: address [TAB] reference
    p2 = doc.add_paragraph()
    _add_right_tab_stop(p2, tab_pos)
    run_addr = p2.add_run(f"{data.client_address},")
    _style_run(run_addr)
    p2.add_run("\t")
    run_ref = p2.add_run(ref_str)
    _style_run(run_ref, bold=True)


def _add_subject(doc: Document, subject_text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Subject: {subject_text}")
    _style_run(run)


def _add_greeting(doc: Document, client_name: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Dear Engr, {client_name}")
    _style_run(run)


def _add_intro(doc: Document) -> None:
    _add_body_text(doc,
        "As requested, please find herewith attached our offer for Fire Alarm System."
    )


def _add_intro_installation(doc: Document) -> None:
    _add_body_text(doc,
        "As requested, please find herewith attached our offer for above mentioned "
        "systems. The Offer is based on the drawings provided. We have proposed "
        "material supply and installation as requested."
    )


def _add_scope(doc: Document, option: int) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("SCOPE")
    _style_run(run, bold=True, underline=True)

    if option == 1:
        _add_body_text(doc,
            "Price includes Supply of equipment mentioned in attached point-schedule, "
            "warranty, programming, testing & commissioning."
        )
    elif option == 2:
        # Partial installation — no conduiting
        _add_body_text(doc,
            "Price includes Supply of equipment mentioned in attached point-schedule, "
            "engineering support which includes preparation of Single Line diagrams, "
            "Installation of devices we supplied, cable pulling, device fixing, "
            "programming, testing and commissioning of equipment we supplied, Client Staff "
            "training, O&M Manuals and Warranty support."
        )
    else:  # option 3
        # Full installation — with conduiting
        _add_body_text(doc,
            "Price includes Supply of equipment mentioned in attached point-schedule, "
            "engineering support which includes preparation of Single Line diagrams, "
            "Installation of devices we supplied, conduiting, cable pulling, device fixing, "
            "programming, testing and commissioning of equipment we supplied, Client Staff "
            "training, O&M Manuals and Warranty support."
        )


def _add_warranty(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("WARRANTY:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "Items supplied by us shall be covered under our standard warranty clause that "
        "covers against any material defect or malfunctioning, for a period of 18 months "
        "from date of delivery. Product shall be used as intended. Misuse or wrong application "
        "will not be covered under warranty. We also hope that project maintenance will be "
        "given to us as a separate contract so that we can maintain the system in a proper way."
    )
    _add_body_text(doc,
        "However our warranty coverage shall not include wear & tear, consumables, "
        "abuse/ misuse/ wrong use of components"
    )


def _add_warranty_short(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("WARRANTY:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "However, our warranty coverage shall not include wear & tear, consumables, "
        "abuse/ misuse/ wrong use of components."
    )


def _add_notes_exclusions(doc: Document, data: QuotationData) -> None:
    from .inclusions import build_document_items

    items = build_document_items(
        data.service_option,
        data.inclusion_answers or {},
    )

    if not items:
        return

    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("NOTES & EXCLUSIONS:")
    _style_run(run, bold=True, underline=True)

    col_num_w = Cm(0.8)
    table_width = Emu(int(_CONTENT_WIDTH * 0.85))  # 85% of content width
    col_text_w = Emu(int(table_width) - int(col_num_w))

    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Force fixed layout so Word doesn't auto-resize columns
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    # Add cell padding for readability
    cell_mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", "40"), ("bottom", "40"), ("left", "80"), ("right", "80")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)

    # Set grid column widths at table level
    tbl_grid = table._tbl.tblGrid
    grid_cols = tbl_grid.findall(qn("w:gridCol"))
    grid_cols[0].set(qn("w:w"), str(int(col_num_w)))
    grid_cols[1].set(qn("w:w"), str(int(col_text_w)))

    for i, item in enumerate(items):
        row = table.rows[i]
        row.cells[0].width = col_num_w
        row.cells[1].width = col_text_w

        # Number cell — centered vertically
        cell_num = row.cells[0]
        cell_num.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_num = cell_num.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_num = p_num.add_run(str(i + 1))
        run_num.font.name = _FONT_NAME
        run_num.font.size = Pt(9)

        # Text cell — centered vertically
        cell_text = row.cells[1]
        cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_text = cell_text.paragraphs[0]
        run_text = p_text.add_run(item)
        run_text.font.name = _FONT_NAME
        run_text.font.size = Pt(9)


def _add_cancellation(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("CANCELLATION:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "In case of cancellation of order for whatsoever reasons RGM reserves the right "
        "to charge the purchaser for the cost of such cancellations in accordance with "
        "the actual stage of processing."
    )


def _add_limitation_of_liability(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("LIMITATION OF LIABILITY:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "The supplier shall not be liable, whether in contract, warranty, failure of remedy "
        "to achieve its essential purpose, tort (including negligence or strict liability) "
        "indemnity, or any other legal or equitable theory for damage to or loss of other "
        "property or equipment, business interruption or lost revenue, profits or sales, "
        "cost of capital, or for any special, incidental, punitive, indirect or consequential "
        "damages or for any other loss, costs or expenses of similar type."
    )
    _add_body_text(doc,
        "The liability of the supplier for any act or omission, product sold, serviced or "
        "furnished directly or indirectly under this agreement, whether in contract, warranty "
        "failure or a remedy to achieve its essential purpose, tort (including negligence or "
        "strict liability) indemnity, or any other legal or equitable theory, will in no event "
        "exceed 1% of the contract value."
    )
    _add_body_text(doc,
        "The rights and remedies contained in this agreement are exclusive, and the parties "
        "accept these remedies in lieu of all other rights and remedies available at law or "
        "otherwise, in contract (including warranty) or in tort (including negligence), for "
        "any and all claims of any nature arising under this agreement or any performance or "
        "breach arising out of this agreement."
    )


def _add_prices_and_terms(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("PRICES AND TERMS OF PAYMENT:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "Without prejudice to any further rights, we may suspend and/ or refuse any supplies "
        "for as long as any due payment remains outstanding for whatsoever reason."
    )
    _add_body_text(doc,
        "Late payments due and payable to supplier shall attract interest at a rated of 12% "
        "per annum accruing from their due date until full settlement of the principal amount. "
        "Payments by the purchaser shall be deemed to be made first against any accrued interest "
        "and then against the outstanding principal amount. The provision of this clause is "
        "without prejudice to any further rights of the supplier in case of payment is delayed "
        "by the purchaser."
    )


def _add_payment_terms(doc: Document, data: QuotationData) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=160, after=60)
    run = p.add_run("Payment terms :")
    _style_run(run, bold=True)

    if data.payment_terms_text:
        for line in data.payment_terms_text.splitlines():
            if line.strip():
                p = doc.add_paragraph()
                run = p.add_run(line)
                _style_run(run)


def _add_time_for_supplies(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run("TIME FOR SUPPLIES:")
    _style_run(run, bold=True, underline=True)

    _add_body_text(doc,
        "Delivery \u2013 14 to 16 weeks from the date of advance payment receipt. "
        "(Partial Delivery against partial payment allowed)"
    )


def _add_validity_and_signature(doc: Document, data: QuotationData) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=100)
    run = p.add_run("Validity \u2013 30 days")
    _style_run(run, bold=True)

    # "Best regards," + name — only if signatory name is provided
    if data.signatory_name:
        _add_body_text(doc, "Best regards,")
        _add_body_text(doc, data.signatory_name)

    # Phone — only if provided (independent of name/signature)
    if data.company_phone:
        _add_body_text(doc, data.company_phone)

    # Signature image — only if uploaded by admin (at the very end)
    if data.signature_bytes:
        p = doc.add_paragraph()
        p.add_run().add_picture(io.BytesIO(data.signature_bytes), width=Cm(3.5))


def _add_product_table(doc: Document, data: QuotationData) -> None:
    # Start product table on a new page
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    has_installation = data.installation_amount > 0
    # header + products + total + vat + grand total (+ optional installation row)
    num_rows = len(data.products) + 4 + (1 if has_installation else 0)
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Force fixed table layout so Word respects our column widths
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    # Add small vertical cell padding for breathing room
    cell_mar = OxmlElement("w:tblCellMar")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), "28")  # ~0.5mm
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)

    # Set grid column widths at the XML level
    col_widths = [_COL_MODEL, _COL_DESC, _COL_QTY, _COL_UNIT, _COL_TOTAL]
    tbl_grid = table._tbl.tblGrid
    grid_cols = tbl_grid.findall(qn("w:gridCol"))
    for gc, w in zip(grid_cols, col_widths):
        gc.set(qn("w:w"), str(int(w)))

    # Also set cell widths per row for full compliance
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    # Header row
    headers = ["Model", "Description", "Qty", " Unit\nPrice ", " Total\nPrice "]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        _set_table_cell(cell, h, bold=True)
        _shade_cell(cell, "D9D9D9")

    # Product rows
    for ri, prod in enumerate(data.products, 1):
        _set_table_cell(table.cell(ri, 0), prod.code)
        _set_table_cell(table.cell(ri, 1), prod.description)
        _set_table_cell(table.cell(ri, 2), _format_qty(prod.quantity),
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_table_cell(table.cell(ri, 3), _format_price(prod.unit_price),
                        align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_table_cell(table.cell(ri, 4), _format_price(prod.total_price),
                        align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Total row
    total_row = len(data.products) + 1
    _set_table_cell(table.cell(total_row, 1), "TOTAL IN SAR", bold=True)
    _set_table_cell(table.cell(total_row, 4), _format_price(data.subtotal),
                    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Installation services row (options 2/3 only)
    next_row = total_row + 1
    if has_installation:
        install_row = next_row
        rate = data.installation_amount / data.device_count if data.device_count else 0
        _set_table_cell(table.cell(install_row, 1), "Installation Services", bold=True)
        _set_table_cell(table.cell(install_row, 2), str(data.device_count),
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_table_cell(table.cell(install_row, 3), _format_price(rate),
                        align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_table_cell(table.cell(install_row, 4), _format_price(data.installation_amount),
                        bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        next_row = install_row + 1

    # VAT row
    vat_row = next_row
    _set_table_cell(table.cell(vat_row, 1), "VAT", bold=True)
    _set_table_cell(table.cell(vat_row, 4), _format_price(data.vat),
                    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Grand total row
    gt_row = vat_row + 1
    _set_table_cell(table.cell(gt_row, 1), "GRAND TOTAL IN SAR", bold=True)
    _set_table_cell(table.cell(gt_row, 4), _format_price(data.grand_total),
                    bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_body_text(doc: Document, text: str, space_after: int = 120) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after / 20)  # twips to pt
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _style_run(run)


def _add_empty_para(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("")
    _style_run(run)


def _set_para_spacing(paragraph, before: int = 0, after: int = 120) -> None:
    """Set paragraph spacing in twips. 120 twips ≈ 6pt — standard professional gap."""
    paragraph.paragraph_format.space_before = Pt(before / 20)
    paragraph.paragraph_format.space_after = Pt(after / 20)


def _style_run(
    run,
    bold: bool = False,
    underline: bool = False,
    size=None,
) -> None:
    run.font.name = _FONT_NAME
    run.font.size = size or _FONT_SIZE
    run.font.color.rgb = _FONT_COLOR
    run.font.bold = bold
    run.font.underline = underline


def _set_date_cell(cell: _Cell, day: int, suffix: str, month: str, year: int) -> None:
    """Render 'Date: 29th Mar 2026' with superscript ordinal suffix."""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run1 = p.add_run(f"Date: {day}")
    _style_run(run1, bold=True)

    run_sup = p.add_run(suffix)
    _style_run(run_sup, bold=True, size=Pt(7))
    run_sup.font.superscript = True

    run2 = p.add_run(f" {month} {year}")
    _style_run(run2, bold=True)


def _set_cell_text(
    cell: _Cell,
    text: str,
    bold: bool = False,
    align=None,
) -> None:
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    _style_run(run, bold=bold)


def _set_table_cell(
    cell: _Cell,
    text: str,
    bold: bool = False,
    align=None,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    # Clear existing
    for run in p.runs:
        run.text = ""
    run = p.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = Pt(9)
    run.font.bold = bold


def _shade_cell(cell: _Cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_right_tab_stop(paragraph, position) -> None:
    """Add a right-aligned tab stop at the given position (EMU)."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position / 635)))  # EMU to twips
    tabs.append(tab)
    pPr.append(tabs)


def _remove_cell_borders(cell: _Cell) -> None:
    """Remove all borders from a single table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tblPr.append(borders)


def _ordinal_suffix(day: int) -> str:
    if 11 <= day <= 13:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def _format_price(value: float) -> str:
    return f"{value:,.2f}"


def _format_qty(value: int | float) -> str:
    if isinstance(value, float) and value == int(value):
        return str(int(value))
    return str(value)


def _add_anchored_image(
    run, image_path: str,
    cx: int, cy: int,
    pos_h: int, pos_v: int,
    part=None,
) -> None:
    """Add a floating (anchored) image to a run with absolute positioning in EMU."""
    # Add image part to the document
    r_element = run._r
    if part is None:
        part = r_element.part
    rId, _image = part.get_or_add_image(image_path)

    # Build the anchor XML matching the reference template structure
    anchor_xml = (
        f'<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        f' behindDoc="1" distT="0" distB="0" distL="114935" distR="114935"'
        f' simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="0">'
        f'  <wp:simplePos x="0" y="0"/>'
        f'  <wp:positionH relativeFrom="column">'
        f'    <wp:posOffset>{pos_h}</wp:posOffset>'
        f'  </wp:positionH>'
        f'  <wp:positionV relativeFrom="paragraph">'
        f'    <wp:posOffset>{pos_v}</wp:posOffset>'
        f'  </wp:positionV>'
        f'  <wp:extent cx="{cx}" cy="{cy}"/>'
        f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'  <wp:wrapNone/>'
        f'  <wp:docPr id="0" name=""/>'
        f'  <wp:cNvGraphicFramePr>'
        f'    <a:graphicFrameLocks noChangeAspect="1"/>'
        f'  </wp:cNvGraphicFramePr>'
        f'  <a:graphic>'
        f'    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'      <pic:pic>'
        f'        <pic:nvPicPr>'
        f'          <pic:cNvPr id="0" name=""/>'
        f'          <pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>'
        f'        </pic:nvPicPr>'
        f'        <pic:blipFill>'
        f'          <a:blip r:embed="{rId}"/>'
        f'          <a:stretch><a:fillRect/></a:stretch>'
        f'        </pic:blipFill>'
        f'        <pic:spPr>'
        f'          <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'          <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'        </pic:spPr>'
        f'      </pic:pic>'
        f'    </a:graphicData>'
        f'  </a:graphic>'
        f'</wp:anchor>'
    )

    from lxml import etree
    drawing = OxmlElement("w:drawing")
    anchor_el = etree.fromstring(anchor_xml)
    drawing.append(anchor_el)
    r_element.append(drawing)
